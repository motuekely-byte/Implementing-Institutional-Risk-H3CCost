!pip install -q python-docx openpyxl

import pandas as pd
import numpy as np
import csv, os, io
from google.colab import files
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.metrics import accuracy_score, roc_auc_score
from sklearn.neural_network import MLPClassifier
from docx import Document
from docx.shared import Inches

# 🧩 Utils
def to_num(s):
    """Conversion robuste en float (virgules, espaces…)."""
    return pd.to_numeric(
        pd.Series(s).astype(str)
          .str.replace(',', '.', regex=False)
          .str.replace(r'\s+', '', regex=True),
        errors='coerce'
    )

def resample_with_weights(X, y, weights):
    """Rééchantillonnage basé sur les weights pour simuler sample_weight."""
    n_samples = len(y)
    if len(weights) != n_samples:
        raise ValueError(f"Length of weights ({len(weights)}) does not match number of samples ({n_samples})")
    total_weight = np.sum(weights)
    if total_weight == 0:
        weights = np.ones(n_samples) / n_samples
    else:
        weights = weights / total_weight
    indices = np.random.choice(np.arange(n_samples), size=n_samples, p=weights, replace=True)
    return X[indices], y[indices]

def standardize_target(y_series, dataset_name):
    """
    Mappe la cible en binaire :
      - IMPAYE / DEFAULT / 0  -> 0
      - PAYE / NON_DEFAULT / 1 -> 1
    """
    print(f"📋 Valeurs uniques dans target pour {dataset_name}: {y_series.unique()}")
    y = y_series.astype(str).str.strip().str.upper()
    y = y.replace({
        'IMPAYE': 0, 'DEFAULT': 0, 'DEF': 0, '0': 0, 'NO': 0, 'N': 0, 'FALSE': 0, 'NEGATIVE': 0,
        'PAYE': 1, 'NON_DEFAULT': 1, '1': 1, 'YES': 1, 'Y': 1, 'TRUE': 1, 'POSITIVE': 1,
        'APPROVED': 1, 'ACCEPT': 1, 'PAID': 1
    }, inplace=False)
    y = pd.to_numeric(y, errors='coerce')
    nans = y.isna().sum()
    if nans > 0:
        print(f"⚠️ {nans} valeurs non reconnues dans target pour {dataset_name}, remplacées par NaN")
        invalid_values = y_series[y.isna()].unique()
        print(f"📋 Valeurs non reconnues: {invalid_values}")
    if y.isna().all():
        raw = y_series.astype(str).str.lower()
        if set(raw.unique()) <= {'0', '1'}:
            y = raw.astype(int)
            print(f"✅ Conversion directe 0/1 pour {dataset_name}")
        else:
            raise ValueError(f"[{dataset_name}] Toutes les valeurs de target sont non reconnues: {invalid_values}")
    return y.fillna(0).astype(np.int64)

# 💰 Fonctions de coût
def misclassification_costs(y_true, y_pred, amount, rate, duration, lgd=0.6):
    """MCI: FP = m*t*(d/1200), FN = m*lgd, correct=0."""
    costs = []
    for yt, yp, m, t, d in zip(y_true, y_pred, amount, rate, duration):
        if yt == yp:
            c = 0.0
        elif yt == 1 and yp == 0:      # FP (vrai PAYE refusé)
            c = m * t * (d / 1200.0)
        elif yt == 0 and yp == 1:      # FN (vrai IMPAYE accepté)
            c = m * lgd
        else:
            c = 0.0
        costs.append(c)
    return np.array(costs, dtype=float)

def compute_h3c(y_true, y_pred, amt, rate, dur, gamma=0.07, lgd=0.6):
    """Version simplifiée fournie par l'utilisateur."""
    h3c_list = []
    for yt, yp, m, t, d in zip(y_true, y_pred, amt, rate, dur):
        if yt == yp:
            h3c = 0.0
        else:
            if yt == 1 and yp == 0:      # FP
                c_m = m * t * (d / 1200.0)
            elif yt == 0 and yp == 1:    # FN
                c_m = m * lgd
            else:
                c_m = 0.0
            c_s = 0.01 * m  # coût stratégique simplifié
            c_i = gamma * m # coût incitatif
            h3c = 0.7 * c_m + 0.2 * c_s + 0.1 * c_i
        h3c_list.append(h3c)
    return np.array(h3c_list, dtype=float)

# 🧪 Entraînement/évaluation
def run_models_on_df(df_raw, dataset_name, default_rate=0.10, default_dur=36, default_lgd=0.5, test_size=0.3, random_state=42):
    df = df_raw.copy()
    mapping = {
        "datafb.csv": {"amount": "Montant", "rate": "Taux", "duration": "Durée", "target": "ENIMPAYEOUPAS", "lgd": 0.60},
        "Portuguese.csv": {"amount": "balance", "rate": None, "duration": "duration", "target": "y", "lgd": 0.45},
        "Loan Dataset Kaggle.csv": {"amount": "Loan_Amount_Requested", "rate": "Interest_Rate", "duration": "Loan_Term", "target": "Loan_Approval_Status", "lgd": 0.50},
        "german_credit_data.csv": {"amount": "Credit amount", "rate": None, "duration": "Duration", "target": None, "lgd": 0.25},
        "credit card taiwan.csv": {"amount": "LIMIT_BAL", "rate": None, "duration": "PAY_0", "target": "default.payment.next.month", "lgd": 0.40},
        "australian.csv": {"amount": "A2", "rate": "A3", "duration": "A4", "target": "CLASS", "lgd": 0.375}
    }
    lgd = mapping.get(dataset_name, {}).get("lgd", default_lgd)

    # Renommage des colonnes selon column_mapping
    renamed = {}
    for std_name, col_name in mapping.get(dataset_name, {}).items():
        if col_name and col_name in df.columns:
            renamed[col_name] = std_name
    df.rename(columns=renamed, inplace=True)

    # Injecter valeurs par défaut si colonnes absentes
    if 'amount' not in df.columns:
        df['amount'] = 5000.0
        print(f"⚠️ Montant absent dans {dataset_name}, valeur par défaut 5000 injectée.")
    if 'rate' not in df.columns:
        df['rate'] = default_rate
        print(f"⚠️ Taux absent dans {dataset_name}, valeur par défaut {default_rate*100}% injectée.")
    if 'duration' not in df.columns:
        df['duration'] = default_dur
        print(f"⚠️ Durée absente dans {dataset_name}, valeur par défaut {default_dur} mois injectée.")
    if 'target' not in df.columns:
        print(f"⏭️ Skip du dataset '{dataset_name}' car la colonne cible est absente.")
        return None, None

    # Numériser les colonnes clés et vérifier leur validité
    amount = to_num(df['amount']).fillna(5000.0).astype(float)
    rate = to_num(df['rate']).fillna(default_rate).astype(float)
    duration = to_num(df['duration']).fillna(default_dur).astype(float)

    print(f"📋 Stats pour {dataset_name} - amount: min={amount.min()}, max={amount.max()}, NaN={amount.isna().sum()}")
    print(f"📋 Stats pour {dataset_name} - rate: min={rate.min()}, max={rate.max()}, NaN={rate.isna().sum()}")
    print(f"📋 Stats pour {dataset_name} - duration: min={duration.min()}, max={duration.max()}, NaN={duration.isna().sum()}")

    amount = np.where((amount <= 0) | (amount.isna()), 5000.0, amount)
    rate = np.where((rate <= 0) | (rate.isna()), default_rate, rate)
    duration = np.where((duration <= 0) | (duration.isna()), default_dur, duration)

    # Cible binaire
    y = standardize_target(df['target'], dataset_name)

    # Features encodées
    X = df.drop(columns=['target'])
    if 'ID' in X.columns:
        X = X.drop(columns=['ID'])
    for col in X.columns:
        if X[col].dtype == 'object':
            X[col] = LabelEncoder().fit_transform(X[col].astype(str))
    X_enc = X.fillna(0).astype(np.float32)

    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X_enc).astype(np.float32)

    # Split
    X_train, X_test, y_train, y_test, amt_tr, amt_te, rate_tr, rate_te, dur_tr, dur_te = train_test_split(
        X_scaled, y, amount, rate, duration,
        test_size=test_size, random_state=random_state, stratify=y if len(np.unique(y))==2 else None
    )

    # Convertir en arrays numpy pour éviter les problèmes d'indexation
    X_train = np.asarray(X_train)
    y_train = np.asarray(y_train)
    X_test = np.asarray(X_test)
    y_test = np.asarray(y_test)

    # Poids d'entraînement
    mci_train = np.where(
        y_train == 1,
        amt_tr * rate_tr * (dur_tr / 1200.0),
        amt_tr * lgd
    ).astype(float)
    mci_train = np.where(mci_train <= 0, 1.0, mci_train)
    mci_train = mci_train / (mci_train.mean() + 1e-12)

    h3c_tmp = compute_h3c(y_train, 1 - y_train, amt_tr, rate_tr, dur_tr, lgd=lgd)
    h3c_train = np.where(h3c_tmp <= 0, 1.0, h3c_tmp)
    h3c_train = h3c_train / (h3c_train.mean() + 1e-12)

    # Modèles (MLP avec rééchantillonnage)
    common_params = dict(hidden_layer_sizes=(100,), max_iter=2000, random_state=random_state)
    models = {
        "MLP CLASSIC": MLPClassifier(**common_params),
        "MLP CS": MLPClassifier(**common_params),
        "MLP H3C": MLPClassifier(**common_params)
    }

    per_row = []
    summary_rows = []
    predictions = {}
    for name, model in models.items():
        # Entraînement avec rééchantillonnage
        if name == "MLP CLASSIC":
            X_train_used, y_train_used = X_train, y_train
        elif name == "MLP CS":
            X_train_used, y_train_used = resample_with_weights(X_train, y_train, mci_train)
        elif name == "MLP H3C":
            X_train_used, y_train_used = resample_with_weights(X_train, y_train, h3c_train)

        model.fit(X_train_used, y_train_used)

        y_pred = model.predict(X_test)
        y_proba = model.predict_proba(X_test)[:, 1]

        # Coûts
        cost_mci = misclassification_costs(y_test, y_pred, amt_te, rate_te, dur_te, lgd=lgd)
        cost_h3c = compute_h3c(y_test, y_pred, amt_te, rate_te, dur_te, lgd=lgd)

        # Baseline
        baseline_pred = np.ones_like(y_test)
        baseline_mci = misclassification_costs(y_test, baseline_pred, amt_te, rate_te, dur_te, lgd=lgd)
        baseline_h3c = compute_h3c(y_test, baseline_pred, amt_te, rate_te, dur_te, lgd=lgd)

        resid_cost = cost_h3c.sum() if "H3C" in name else cost_mci.sum()
        base_cost = baseline_h3c.sum() if "H3C" in name else baseline_mci.sum()
        gain = base_cost - resid_cost
        prop_gain = (gain / base_cost) if base_cost > 0 else 0.0

        # Résumé
        summary_rows.append({
            "Dataset": dataset_name,
            "Type": "privé" if dataset_name == "datafb.csv" else "public",
            "Modèle": name,
            "Accuracy": round(accuracy_score(y_test, y_pred)*100, 2),
            "ROC AUC": round(roc_auc_score(y_test, y_proba)*100, 2) if len(np.unique(y_test))==2 else np.nan,
            "Coût Moyen": round(resid_cost / max(len(y_test), 1), 2),
            "Coût Global résiduel": round(resid_cost, 2),
            "Gain lié à l'utilisation du modele": round(gain, 2),
            "proportion de gain": f"{(prop_gain*100):.2f}%"
        })

        # Stocker les prédictions et coûts
        predictions[name] = {
            "y_pred": y_pred,
            "cost_mci": cost_mci,
            "cost_h3c": cost_h3c
        }

    # Construire le DataFrame des résultats détaillés
    results_df = pd.DataFrame({
        'amount': amt_te,
        'rate': rate_te,
        'duration': dur_te,
        'y_true': y_test,
        'y_pred_classic': predictions["MLP CLASSIC"]["y_pred"],
        'cost_mci_classic': predictions["MLP CLASSIC"]["cost_mci"],
        'y_pred_cs': predictions["MLP CS"]["y_pred"],
        'cost_mci_cs': predictions["MLP CS"]["cost_mci"],
        'y_pred_h3c': predictions["MLP H3C"]["y_pred"],
        'cost_h3c': predictions["MLP H3C"]["cost_h3c"]
    })

    return results_df, summary_rows

# 📥 Chargement des données
print("Sélectionne tes fichiers CSV…")
uploaded = files.upload()

datasets = {}
encodings = ['latin1', 'ISO-8859-1', 'utf-8']
seps = [';', ',', '\t']
dataset_separators = {
    "datafb.csv": ";",
    "Portuguese.csv": ";",
    "Loan Dataset Kaggle.csv": ",",
    "german_credit_data.csv": ",",
    "credit card taiwan.csv": ",",
    "australian.csv": ";"
}

for fname, content in uploaded.items():
    df = None
    sep = dataset_separators.get(fname, ",")
    for enc in encodings:
        try:
            df = pd.read_csv(io.BytesIO(content), encoding=enc, sep=sep, engine='python', quoting=csv.QUOTE_MINIMAL)
            if df.shape[1] == 1:
                continue
            print(f"✅ {fname} chargé (enc={enc}, sep='{sep}', shape={df.shape})")
            break
        except Exception:
            continue
    if df is None:
        for enc in encodings:
            for alt_sep in seps:
                if alt_sep == sep:
                    continue
                try:
                    df = pd.read_csv(io.BytesIO(content), encoding=enc, sep=alt_sep, engine='python', quoting=csv.QUOTE_MINIMAL)
                    if df.shape[1] == 1:
                        continue
                    print(f"✅ {fname} chargé (enc={enc}, sep='{alt_sep}', shape={df.shape})")
                    break
                except Exception:
                    continue
            if df is not None:
                break
    if df is None:
        print(f"❌ Échec de lecture pour {fname}")
    else:
        datasets[fname] = df

if not datasets:
    raise SystemExit("Aucun fichier lisible.")

# ▶️ Exécution sur chaque dataset
all_summaries = []
per_dataset_results = {}
for fname, df in datasets.items():
    try:
        res_df, summary_rows = run_models_on_df(df, dataset_name=fname)
        if res_df is not None and summary_rows:
            per_dataset_results[fname] = res_df
            all_summaries.extend(summary_rows)
    except Exception as e:
        print(f"⚠️ Skip {fname} : {e}")

if not all_summaries:
    raise SystemExit("Aucun résultat produit (colonnes cibles absentes ?).")

# 💾 Exports
# 1) Excel
excel_path = "resultats_par_jeu_mlp.xlsx"
with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
    for fname, df in per_dataset_results.items():
        sheet = os.path.splitext(os.path.basename(fname))[0][:31] or "Feuille"
        df.to_excel(writer, index=False, sheet_name=sheet)
print(f"📘 Excel écrit: {excel_path}")

# 2) CSV global
summary_df = pd.DataFrame(all_summaries, columns=[
    "Dataset", "Type", "Modèle", "Accuracy", "ROC AUC", "Coût Moyen",
    "Coût Global résiduel", "Gain lié à l'utilisation du modele", "proportion de gain"
])
csv_path = "resume_global_mlp.csv"
summary_df.to_csv(csv_path, index=False, sep=';')
print(f"📄 CSV global écrit: {csv_path}")

# 3) Rapport Word
doc = Document()
doc.add_heading("Rapport — Évaluation MLP Credit Scoring", 0)

doc.add_heading("1. Introduction", level=1)
doc.add_paragraph(
    "Ce rapport évalue trois modèles MLP (Classic, CS, H3C) sur différents datasets de credit scoring. "
    "Les colonnes sont harmonisées selon un mapping spécifique. "
    "Les coûts sont calculés via MCI et H3C, avec un baseline 'accepter tout'."
)

doc.add_heading("2. Méthodologie", level=1)
doc.add_paragraph(
    "Colonnes renommées via column_mapping. Valeurs par défaut : amount=5000, rate=0.10, duration=36. "
    "Cible standardisée en binaire (0=IMPAYE, 1=PAYE). MCI : FP = montant×taux×(durée/1200), FN = montant×LGD. "
    "H3C : 0.7×coût_mci + 0.2×(0.01×montant) + 0.1×(γ×montant). Modèles entraînés avec rééchantillonnage pour CS et H3C."
)

doc.add_heading("3. Jeux de données", level=1)
doc.add_paragraph(f"Datasets traités : {', '.join(datasets.keys())}")

doc.add_heading("4. Résultats globaux", level=1)
t = doc.add_table(rows=1, cols=9)
t.style = 'Table Grid'
t.rows[0].cells[0].text = "Dataset"
t.rows[0].cells[1].text = "Type"
t.rows[0].cells[2].text = "Modèle"
t.rows[0].cells[3].text = "Accuracy"
t.rows[0].cells[4].text = "ROC AUC"
t.rows[0].cells[5].text = "Coût Moyen"
t.rows[0].cells[6].text = "Coût Global résiduel"
t.rows[0].cells[7].text = "Gain lié à l'utilisation du modele"
t.rows[0].cells[8].text = "proportion de gain"
for _, r in summary_df.iterrows():
    row = t.add_row().cells
    row[0].text = str(r["Dataset"])
    row[1].text = str(r["Type"])
    row[2].text = str(r["Modèle"])
    row[3].text = f'{r["Accuracy"]}'
    row[4].text = f'{r["ROC AUC"]}' if not pd.isna(r["ROC AUC"]) else 'N/A'
    row[5].text = f'{r["Coût Moyen"]}'
    row[6].text = f'{r["Coût Global résiduel"]}'
    row[7].text = f'{r["Gain lié à l\'utilisation du modele"]}'
    row[8].text = str(r["proportion de gain"])

doc.add_heading("5. Interprétation", level=1)
doc.add_paragraph(
    "Le gain est calculé par rapport à une baseline 'accepter tout'. "
    "MLP H3C et CS optimisent les coûts via rééchantillonnage, mais l'accuracy peut varier. "
    "Les datasets avec cibles manquantes ou valeurs invalides peuvent être ignorés."
)

doc.add_heading("6. Conclusion", level=1)
doc.add_paragraph(
    "Cette pipeline traite plusieurs datasets avec harmonisation des colonnes et rééchantillonnage. "
    "Optimisations futures : réglage des hyperparamètres, gestion des valeurs manquantes dans target."
)

word_path = "rapport_credit_scoring_mlp.docx"
doc.save(word_path)
print(f"📝 Rapport Word écrit: {word_path}")

# Téléchargements
files.download(excel_path)
files.download(csv_path)
files.download(word_path)